from docx import Document
from docx.shared import Pt, Cm
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn

# 复制整张表
def copy_table(source_table, target_document):
    # 在目标文档中创建一个新的表格，与源表格大小一致
    target_table = target_document.add_table(rows=len(source_table.rows), cols=len(source_table.columns))
    
    # 复制表格属性
    target_table.style = 'Table Grid' #source_table.style
    target_table.alignment = source_table.alignment
    target_table.autofit = source_table.autofit

    # 遍历源表格的行和列，复制单元格的内容和属性
    for row_idx, source_row in enumerate(source_table.rows):
        for col_idx, source_cell in enumerate(source_row.cells):
            target_cell = target_table.cell(row_idx, col_idx)
            # 复制单元格内容
            target_cell.text = source_cell.text

            # 复制单元格的段落格式
            for p_idx, source_paragraph in enumerate(source_cell.paragraphs):
                target_paragraph = target_cell.paragraphs[p_idx]
                target_paragraph.style = source_paragraph.style
                target_paragraph.alignment = source_paragraph.alignment

                for source_run in source_paragraph.runs:
                    target_run = target_paragraph.add_run(source_run.text)
                    # 复制字体格式
                    target_run.bold = source_run.bold
                    target_run.italic = source_run.italic
                    target_run.underline = source_run.underline
                    if source_run.font.size:
                        target_run.font.size = source_run.font.size
                    if source_run.font.name:
                        target_run.font.name = source_run.font.name
                    if source_run.font.color.rgb:
                        target_run.font.color.rgb = source_run.font.color.rgb

            # 复制单元格格式
            copyCellFormat(source_cell, target_cell)

            # 复制单元格合并情况
            if source_cell._tc != target_cell._tc:  # 避免重复合并
                if source_cell._element.grid_span is not None:
                    print(row_idx, col_idx, source_cell._element.grid_span)
                    target_cell.merge(target_table.cell(row_idx, col_idx + source_cell._element.grid_span - 1))

    set_table_borders(target_table)
    return target_table

# 复制指定列
def copy_column_format_and_content(src_table, dest_table, src_index, dest_index):
    """
    将源表格中的某列内容和格式复制到目标表格的指定列。
    :param src_table: 源表格对象
    :param dest_table: 目标表格对象
    :param src_index: 源列索引
    :param dest_index: 要复制的列索引
    """
    if len(src_table.rows) != len(dest_table.rows):
        print(len(src_table.rows), len(dest_table.rows), ' : rows is not same')

    # 复制表格属性
    dest_table.style = 'Table Grid' #source_table.style
    dest_table.alignment = src_table.alignment
    dest_table.autofit = src_table.autofit

    for i, source_cell in enumerate(src_table.columns[src_index].cells):

        # 获取目标表格中的单元格
        target_cell = dest_table.cell(i, dest_index)
        
        # 复制单元格内容
        #target_cell.text = source_cell.text

        # 复制单元格的段落格式
        for p_idx, source_paragraph in enumerate(source_cell.paragraphs):
            target_paragraph = target_cell.paragraphs[p_idx]
            target_paragraph.style = source_paragraph.style
            target_paragraph.alignment = source_paragraph.alignment

            for source_run in source_paragraph.runs:
                target_run = target_paragraph.add_run(source_run.text)
                # 复制字体格式
                target_run.bold = source_run.bold
                target_run.italic = source_run.italic
                target_run.underline = source_run.underline
                if source_run.font.size:
                    target_run.font.size = source_run.font.size
                if source_run.font.name:
                    target_run.font.name = source_run.font.name
                if source_run.font.color.rgb:
                    target_run.font.color.rgb = source_run.font.color.rgb

        # 复制单元格格式
        copy_size = False if dest_index != 0 else True
        copyCellFormat(source_cell, target_cell, copy_size)
        
        # 复制单元格合并情况
        # if source_cell._tc != target_cell._tc:  # 避免重复合并
        #     if source_cell._element.grid_span is not None:
        #         print(i, dest_index, source_cell._element.grid_span)
        #         target_cell.merge(target_table.cell(i, dest_index + source_cell._element.grid_span - 1))

def copy_column_merge(source_table, target_table, src_index, dest_index):
    """
    复制表格中指定列的单元格合并信息。
    :param source_table: 源表格
    :param target_table: 目标表格
    :param src_index: 源列索引
    :param dest_index: 要复制的列索引
    """
    row_count = len(source_table.rows)
    current_merge_start = None

    for row_index in range(row_count):
        source_cell = source_table.cell(row_index, src_index)

        # 检查当前单元格是否为合并区域的起点或范围中的单元格
        v_merge = source_cell._tc.xpath('.//w:vMerge')
        if v_merge:
            v_merge_val = v_merge[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
            if v_merge_val == 'restart':  # 合并的起始单元格
                if current_merge_start is not None:
                    # 在目标表格中执行合并
                    target_table.cell(current_merge_start, dest_index).merge(
                        target_table.cell(row_index - 1, dest_index)
                    )
                current_merge_start = row_index
            elif v_merge_val is None:  # 合并范围中的单元格
                continue
        else:
            # 如果当前不在合并范围，结束上一次的合并
            if current_merge_start is not None:
                target_table.cell(current_merge_start, dest_index).merge(
                    target_table.cell(row_index - 1, dest_index)
                )
                current_merge_start = None

    # 如果最后一段也是合并范围，处理结束
    if current_merge_start is not None:
        target_table.cell(current_merge_start, dest_index).merge(
            target_table.cell(row_count - 1, dest_index)
        )

def copyCellFormat(source_cell, dest_cell, copy_size=True):
    # copy单元格文字方向
    direct = get_text_direction(source_cell)
    set_text_direction(dest_cell, direct)

    # copy单元格大小
    if copy_size:
        w, h = get_cell_size(source_cell)
        set_cell_size(dest_cell, w, h)

    # copy单元格背景色
    color = get_cell_background(source_cell)
    set_cell_background(dest_cell, color)

    # copy单元格垂直对齐格式
    alignment = get_cell_vertical_alignment(source_cell)
    set_cell_vertical_alignment(dest_cell, alignment)
    # copy单元格边框格式
    # borders_info = get_cell_borders(source_cell)
    # print(borders_info)
    # set_cell_borders(dest_cell, borders_info)


def get_text_direction(cell):
    """
    获取单元格文字方向。
    参数:
        cell: 目标单元格对象
    返回:
        文字方向 (字符串)，如 'btLr', 'tbRl', 'lrTb'，若无方向设置则返回 None。
    """
    # 命名空间映射
    namespace = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    # 获取单元格属性
    tcPr = cell._tc.get_or_add_tcPr()
    # 查找文字方向属性
    text_direction = tcPr.find('w:textDirection', namespace)
    if text_direction is not None:
        return text_direction.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
    return None

def set_text_direction(cell, direction="btLr"):
    """
    设置单元格文字方向。
    参数:
        cell: 目标单元格对象
        direction: 文字方向，可选值 'btLr' (从下到上), 'tbRl' (从上到下), 'lrTb' (默认，从左到右)
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    text_direction = OxmlElement('w:textDirection')
    # 使用完整命名空间 URI 设置属性
    text_direction.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', direction)
    tcPr.append(text_direction)

def get_cell_size(cell):
    """
    获取单元格大小。
    参数:
        cell: 目标单元格对象
    返回:
        (宽度, 高度)，单位为 Cm，若无法获取则返回 None。
    """
    # 命名空间映射
    namespace = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    # 获取列宽
    width = None
    gridCol = cell._tc.getparent().getprevious()  # 获取列定义 (w:gridCol)
    if gridCol is not None:
        width_attr = gridCol.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w')
        if width_attr:
            width = int(width_attr) / 567.0  # Twips 转换为 Cm

    # 获取行高
    height = None
    tr = cell._tc.getparent()
    trPr = tr.find('w:trPr', namespace)
    if trPr is not None:
        trHeight = trPr.find('w:trHeight', namespace)
        if trHeight is not None:
            height_attr = trHeight.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
            if height_attr:
                height = int(height_attr) / 567.0  # Twips 转换为 Cm

    return (Cm(width) if width else None, Cm(height) if height else None)

def set_cell_size(cell, width=None, height=None):
    """
    设置单元格的宽度和高度。
    参数:
        cell: 目标单元格对象
        width: 宽度，单位为 Cm（可选）
        height: 高度，单位为 Cm（可选）
    """
    # 设置列宽（宽度无法直接通过单元格设置，需调整列宽）
    if width is not None:
        cell._tc.width = int(width.cm * 567)  # 转换为 twip 单位

    # 设置行高
    if height is not None:
        tr = cell._tc.getparent()
        trPr = tr.get_or_add_trPr()

        # 检查是否已有高度设置
        trHeight = trPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}trHeight')
        if trHeight is None:
            trHeight = OxmlElement('w:trHeight')
            trPr.append(trHeight)

        # 设置高度属性
        trHeight.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', str(int(height.cm * 567)))  # 转换为 twip 单位

def get_cell_background(cell):
    """
    获取单元格背景色。
    参数:
        cell: 单元格对象
    返回:
        背景色的十六进制颜色代码 (如 'FFFFFF')，如果没有设置，则返回 None。
    """
    # 命名空间映射
    namespace = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    # 获取单元格属性
    tcPr = cell._tc.get_or_add_tcPr()

    # 查找背景色属性 <w:shd w:fill="...">
    shd = tcPr.find('w:shd', namespace)
    if shd is not None:
        return shd.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')

    return None

def set_cell_background(cell, color=None):
    """
    设置单元格的背景色。
    参数:
        cell: 单元格对象
        color: 背景色，格式为十六进制颜色代码（如 'FF0000'），或传递 'clear' 来清除背景色。
    """
    # 定义命名空间
    namespace = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    # 获取单元格的属性
    tcPr = cell._tc.get_or_add_tcPr()

    # 查找或创建背景色元素 <w:shd>
    shd = tcPr.find('w:shd', namespace)
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)

    # 如果传入的 color 为 'clear'，则清除背景色
    if color == 'clear':
        shd.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill', '')
    elif color:
        # 设置背景色，color 格式为十六进制字符串，如 'FF0000'（红色）
        shd.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill', color)

def get_cell_borders(cell):
    """
    获取单元格的边框格式（类型、粗细、间距）。
    参数:
        cell: 要获取边框的单元格对象
    返回:
        dict: 包含每个边的边框属性（如类型、粗细、间距）
    """
    # 获取单元格的XML元素
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # 获取 <w:tblBorders> 元素
    tbl_borders = tcPr.find('.//w:tblBorders', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})

    if tbl_borders is None:
        return None

    # 用于存储边框信息的字典
    borders_info = {}

    # 遍历每个边框（top, left, bottom, right, insideH, insideV）
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = tbl_borders.find(f'w:{border_name}', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
        if border is not None:
            borders_info[border_name] = {
                'type': border.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val'),
                'size': border.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sz'),
                'space': border.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}space')
            }

    return borders_info

def set_cell_borders(cell, borders_info = None):
    """
    为单个单元格设置边框。
    参数:
        cell: 要设置边框的单元格对象
    """
    if borders_info is None:
        borders_info = {
            'top': {'type': 'single', 'size': '4', 'space': '0'},
            'left': {'type': 'single', 'size': '4', 'space': '0'},
            'bottom': {'type': 'single', 'size': '4', 'space': '0'},
            'right': {'type': 'single', 'size': '4', 'space': '0'},
            'insideH': {'type': 'single', 'size': '4', 'space': '0'},
            'insideV': {'type': 'single', 'size': '4', 'space': '0'},
        }
    
    # 获取单元格的XML元素
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # 创建 <w:tblBorders> 元素，并设置每个边框
    tbl_borders = OxmlElement('w:tblBorders')

    # 定义命名空间
    namespace = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    # 为每个边（top, left, bottom, right, insideH, insideV）设置边框
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        info = borders_info[border_name]
        border.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', info['type'])  # 设置边框类型为单线
        border.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}space', info['space'])     # 设置边框的空间为 0
        border.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sz', info['size'])        # 设置边框的粗细（大小）
        tbl_borders.append(border)

    # 将边框设置添加到单元格的属性中
    tcPr.append(tbl_borders)

def set_table_borders(table):
    """
    为整个表格设置所有单元格的边框。
    参数:
        table: 要设置边框的表格对象
    """
    for row in table.rows:
        for cell in row.cells:
            set_cell_borders(cell)

def get_cell_vertical_alignment(cell):
    """
    获取单元格的垂直居中属性。
    :param cell: 要检查的单元格对象
    :return: 'center' 如果是垂直居中，否则返回 None
    """
    # 获取单元格的底层 XML
    tc = cell._tc
    tcPr = tc.find('.//w:tcPr', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
    
    if tcPr is not None:
        # 查找 w:vAlign 元素
        v_align = tcPr.find('.//w:vAlign', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
        if v_align is not None:
            # 获取 vAlign 的值
            v_align_value = v_align.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
            if v_align_value == 'center':
                return 'center'
            elif v_align_value == 'top':
                return 'top'
            elif v_align_value == 'bottom':
                return 'bottom'
    return None  # 如果没有垂直居中，则返回 None

def set_cell_vertical_alignment(cell, alignment):
    """
    设置单元格的垂直对齐方式。
    :param cell: 单元格对象
    :param alignment: 垂直对齐方式，'top', 'center' 或 'bottom'
    """
    # 获取单元格的底层 XML
    tc = cell._tc  # 获取单元格的 XML
    tcPr = tc.find('.//w:tcPr', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})

    if tcPr is None:
        # 如果 tcPr 不存在，创建它
        tcPr = OxmlElement('w:tcPr')
        tc.append(tcPr)

    # 创建 w:vAlign 元素
    v_align = OxmlElement('w:vAlign')
    
    # 根据传入的对齐方式设置 w:vAlign 元素的 val 属性
    if alignment == 'top':
        v_align.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'top')
    elif alignment == 'center':
        v_align.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'center')
    elif alignment == 'bottom':
        v_align.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'bottom')
    else:
        raise ValueError("Invalid alignment value. Use 'top', 'center', or 'bottom'.")

    # 将 vAlign 元素添加到 tcPr 中
    tcPr.append(v_align)

def replaceCellContent(cell, new_text):
    """
    替换表格单元格内容，同时保持原有格式。
    :param cell: 要修改的单元格对象
    :param new_text: 替换后的文本内容
    """
    # 清空单元格的所有段落内容
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ""

    # 将新文本写入第一个段落，保留格式
    if cell.paragraphs:
        # 如果段落数量大于1，删除除第一个段落外的所有段落
        paragraphs = cell.paragraphs
        while len(paragraphs) > 1:
            cell._element.remove(paragraphs[1]._element)  # 删除第二个及之后的段落
            paragraphs = cell.paragraphs  # 更新段落列表
        paragraph = cell.paragraphs[0]
        if not paragraph.runs:  # 如果 runs 为空
            new_run = paragraph.add_run(new_text)  # 添加一个新的 Run
            # new_run.bold = True  # 设置字体加粗
            # new_run.italic = True  # 设置字体斜体
            # print("成功添加新文本到空段落中")
        else:
            paragraph.runs[0].text = new_text  # 写入第一个运行
    else:
        # 如果单元格为空，则创建一个段落和运行
        paragraph = cell.add_paragraph()
        paragraph.add_run(new_text)

def mergeCellWithOneValue(start_cell, end_cell):
    text = start_cell.text
    start_cell.merge(end_cell)
    replaceCellContent(start_cell, text)

def set_cell_content_format(cell, text=None, style=None, font_bold=None, font_size=None, font_name=None, font_color=None):
    if cell.paragraphs:
        paragraph = cell.paragraphs[0]
        if not paragraph.runs:  # 如果 runs 为空
            new_run = paragraph.add_run()  # 添加一个新的 Run
    else:
        # 如果单元格为空，则创建一个段落和运行
        paragraph = cell.add_paragraph()
        paragraph.add_run()

    # 确保 rPr 存在
    rPr = cell.paragraphs[0].runs[0]._element.get_or_add_rPr()
    # 确保 rFonts 存在
    if rPr.find(qn('w:rFonts')) is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    else:
        rFonts = rPr.find(qn('w:rFonts'))
    rFonts.set(qn('w:eastAsia'), u'宋体')  # 中文字体
    # cell.paragraphs[0].runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    
    if text:
        paragraph.runs[0].text = str(text)
    if style:
        cell.paragraphs[0].alignment = style
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if font_bold:
        cell.paragraphs[0].runs[0].font.bold = font_bold
    if font_size:    
        cell.paragraphs[0].runs[0].font.size = font_size  # 设置字号
    if font_name:
        cell.paragraphs[0].runs[0].font.name = font_name # "Times New Roman"
    if font_color:
        cell.paragraphs[0].runs[0].font.color.rgb = font_color

def set_cell_content(cell, text):
    text = str(text)
    if cell.paragraphs:
        paragraph = cell.paragraphs[0]
        if not paragraph.runs:  # 如果 runs 为空
            new_run = paragraph.add_run(text)  # 添加一个新的 Run
            # new_run.bold = True  # 设置字体加粗
            # new_run.italic = True  # 设置字体斜体
            # print("成功添加新文本到空段落中")
        else:
            paragraph.runs[0].text = text  # 写入第一个运行
    else:
        # 如果单元格为空，则创建一个段落和运行
        paragraph = cell.add_paragraph()
        paragraph.add_run(text)

def set_table_width(table, width_cm):
    # 获取表格底层 XML
    tbl = table._tbl  # table._tbl 是 CT_Tbl 对象
    
    # 获取或添加 tblPr（表格属性）
    tblPr = tbl.tblPr
    if tblPr is None:  # 如果 tblPr 不存在，则创建
        tblPr = OxmlElement('w:tblPr')
        tbl.append(tblPr)

    # 查找或创建 tblW 元素
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)

    # 设置宽度
    tblW.set(qn('w:type'), 'dxa')
    tblW.set(qn('w:w'), str(int((width_cm / 2.54) * 1440)))  # 宽度单位为 dxa

def add_image_to_cell(table, row_idx, col_idx, image_path, width=None, height=None):
    """
    在指定表格单元格中插入图片。
    
    :param table: 表格对象
    :param row_idx: 行索引
    :param col_idx: 列索引
    :param image_path: 图片路径
    :param width: 图片宽度（可选）
    :param height: 图片高度（可选）
    """
    # 获取指定单元格
    cell = table.cell(row_idx, col_idx)
    
    if cell.paragraphs:
        paragraph = cell.paragraphs[0]
        if not paragraph.runs:  # 如果 runs 为空
            new_run = paragraph.add_run()  # 添加一个新的 Run
    else:
        # 如果单元格为空，则创建一个段落和运行
        paragraph = cell.add_paragraph()
        paragraph.add_run()

    # 在单元格中第一个段落
    paragraph = cell.paragraphs[0]
    
    # 插入图片
    run = paragraph.runs[0]
    if width or height:
        run.add_picture(image_path, width=width, height=height)
    else:
        run.add_picture(image_path)

if __name__ == '__main__':
    # 使用示例
    source_doc = Document("./templates/pow_template1.docx")
    target_doc = Document()

    # 假设我们要复制第一个表格
    # source_table = source_doc.tables[0]
    # copied_table = copy_table(source_table, target_doc)

    source_table = source_doc.tables[0]
    target_table = target_doc.add_table(14, 10)
    # target_table.style = 'Table Grid'
    # target_table.alignment = source_table.alignment
    # target_table.autofit = source_table.autofit
    # copy_column_format_and_content(source_table, target_table, 0, 0)
    # copy_column_format_and_content(source_table, target_table, 5, 5)
    for i in range(0, 6):
        copy_column_format_and_content(source_table, target_table, i, i)
    for i in range(6, 8):
        copy_column_format_and_content(source_table, target_table, 6, i)
    copy_column_format_and_content(source_table, target_table, 8, 8)
    copy_column_format_and_content(source_table, target_table, 23, 9)

    for i in range(0, 4):
        if i%2 == 0:
            mergeCellWithOneValue(target_table.cell(3, i), target_table.cell(4, i))
        mergeCellWithOneValue(target_table.cell(5, i), target_table.cell(6, i))
        mergeCellWithOneValue(target_table.cell(10, i), target_table.cell(12, i))
    for i in range(4, len(target_table.rows[0].cells) - 1):
        mergeCellWithOneValue(target_table.cell(4, i), target_table.cell(5, i))
        mergeCellWithOneValue(target_table.cell(6, i), target_table.cell(7, i))
        if i == 4:
            mergeCellWithOneValue(target_table.cell(9, i), target_table.cell(10, i))
    mergeCellWithOneValue(target_table.cell(0, 9), target_table.cell(12, 9))
    # 保存结果
    target_doc.save("target.docx")
