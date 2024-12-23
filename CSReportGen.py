from docx import Document
from docx.shared import RGBColor, Cm, Inches, Mm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
#from docx.oxml import copy_element
from copy import deepcopy

from matplotlib.figure import Figure
import matplotlib.pyplot as plt
from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas
from ReportGen import RepoortGenBase
from utils.docxOp import *
from utils.Utils import FigureManager

class CSReportGenImpl(RepoortGenBase):
    def __init__(self):
        self.allData = {} # { 工况 ： [该工况下的csData] }
        self.figureManager = FigureManager()
        self.figsizes = { 'trace_size' : [6, 6], 'speed_size' : [5, 3.5]} # inch
        super().__init__()
    
    def genReport(self, dataPath:str, **kwargs):
        self.shipParams = kwargs
        self.logger.debug(f"gen cs report from: {dataPath}")
        for field, value in kwargs.items():
            self.logger.info(f"Field: {field}, Value: {value}")

        voyageData = kwargs['voyageData']
        self.voyageCount = len(voyageData)
        # print('v count ', self.voyageCount)
        # 汇总各个工况的数据
        for i, v in enumerate(voyageData):
            if not v['MELoad'] in self.allData:
                self.allData[v['MELoad']] = []
            dataList = self.allData[v['MELoad']]
            no = len(dataList) + 1
            csData = self.fillCsData(no, i, v)
            dataList.append(csData)
        
        # 绘制各工况轨迹图
        self.genTraceFigures()

        # 绘制航速迹程图
        self.genSpeedFigures()

        return self.figureManager

        

    def fillCsData(self, no, index, voyage_data):
        csData = self.analyzer.handleCsData(voyage_data['dataPath'])
        csData.trialNo = self.shipParams.get('hullNo', 'unknown') + '-0' + str(no)
        csData.MELoad = voyage_data['MELoad']
        csData.depth = voyage_data['depth']
        csData.windSpeed = voyage_data['windSpeed']
        csData.windDirection = voyage_data['windOri']
        csData.period = 600
        csData.seq = index
        return csData

    def genTraceFigures(self):
        width = self.figsizes['trace_size'][0]
        height = self.figsizes['trace_size'][1]
        plt.rcParams['font.family'] = 'Times New Roman, SimSun'
        plt.rcParams['xtick.direction'] = 'in'
        plt.rcParams['ytick.direction'] = 'in'
        # plt.rcParams['axes.labelsize'] = 'large' # 14  # 坐标轴标签
        # plt.rcParams['xtick.labelsize'] = 'large' # 12  # x轴刻度
        # plt.rcParams['ytick.labelsize'] = 'large' # 12  # y轴刻度
        for load, datalist in self.allData.items():
            for i, csData in enumerate(datalist):
                figure_name = 'ME' + load + ' ' + csData.trialNo + '轨迹图'
                figure = self.figureManager.get_figure(figure_name)
                figureCanvas = self.figureManager.get_canvas(figure_name)
                if not figure:
                    figure = Figure(figsize=(width, height))
                    figureCanvas = FigureCanvas(figure)
                    self.figureManager.add_figure(csData.seq, figure_name, figure, figureCanvas)
                    csData.traceFigName = figure_name
                ax = figure.add_subplot(111)

                # ax.plot(csData.df['x_0'], csData.df['y_0'], label='Original Curve', color='blue')
                ax.plot(csData.df['x_0_r'], csData.df['y_0_r'], color='blue')
                ax.set_aspect(0.4, adjustable='box')  # y 轴为 x 轴的 0.4 倍, 实际显示为0.8左右，怀疑多个设置参数相互影响了
                xy_lim = {'left': -1500, 'right': 1500, 'bottom': -1000, 'top': 7000} 
                arrow_start_ratio = [0.68, 0.64] # [x, y]
                arrow_end_ratio = [0.68, 0.83]
                text_ratio = [0.73, 0.72]
                ax.annotate(
                    '',  # 注释文本
                    xy=(arrow_end_ratio[0] * xy_lim['right'], arrow_end_ratio[1] * xy_lim['top']),               # 要标注的点 (x, y)
                    xytext=(arrow_start_ratio[0] * xy_lim['right'], arrow_start_ratio[1] * xy_lim['top']),           # 注释文本的位置
                    arrowprops=dict(facecolor='blue', arrowstyle='->')  # 设置箭头样式
                )
                ax.annotate(
                    str(csData.heading) + '°',  # 注释文本
                    xy=(text_ratio[0] * xy_lim['right'], text_ratio[1] * xy_lim['top'])  # 设置文本起点
                )
                # ax.annotate(
                #     csData.heading,  # 注释文本
                #     xy=(arrow_start_ratio[0] * xycoords=['left'], arrow_start_ratio[1] * xycoords=['right']),               # 要标注的点 (x, y)
                #     xytext=(900, 4000),           # 注释文本的位置
                #     arrowprops=dict(facecolor='blue', arrowstyle='->')  # 设置箭头样式
                # )
                # ax.legend()
                title = 'ME' + load + ' ' + csData.trialNo + ' 轨迹图'
                ax.set_title(title)
                figure.tight_layout()
                # ax.set_xlabel('x')
                # ax.set_ylabel('y')
                ax.set_xlim(xy_lim['left'], xy_lim['right'])
                ax.set_ylim(xy_lim['bottom'], xy_lim['top'])
                ax.spines['bottom'].set_linewidth(1.5)
                ax.spines['top'].set_linewidth(1.5)
                ax.spines['left'].set_linewidth(1.5)
                ax.spines['right'].set_linewidth(1.5)
                ax.grid(True)
                # figure.subplots_adjust(left=0.1, right=0.9, top=0.9, bottom=0.1)

                # 在起点和终点画圈
                ax.scatter(csData.df['x_0_r'].iloc[0], csData.df['y_0_r'].iloc[0], s=25, edgecolors='purple', facecolors='none', linewidths=0.5, label='起点')
                ax.scatter(csData.df['x_0_r'].iloc[-1], csData.df['y_0_r'].iloc[-1], s=25, edgecolors='purple', facecolors='none', linewidths=0.5, label='终点')

                # 在起点和终点画小船
                '''小船放大系数'''
                k_larger = 4.0
                self.drawShip(ax, csData.df['x_0_r'].iloc[0], csData.df['y_0_r'].iloc[0], k_larger)
                self.drawShip(ax, csData.df['x_0_r'].iloc[-1], csData.df['y_0_r'].iloc[-1], k_larger)

                # 保存图片，调试时打开
                figureCanvas.draw()
                path = f'{title}.png'
                figureCanvas.print_png(path)

    def drawShip(self, axes, x_center, y_center, k_larger = 1.0):
        ship_half_width = k_larger * 3
        ship_half_length = k_larger * 25
        ship_point = k_larger * 5
        x = [x_center, x_center+3*k_larger, x_center+ship_half_width, x_center, x_center-ship_half_width, x_center-ship_half_width, x_center, x_center]
        y = [y_center+ship_half_length, y_center+ship_point, y_center-ship_half_length, y_center-ship_half_length, y_center-ship_half_length, y_center+ship_point, y_center+ship_half_length, y_center-ship_half_length]
        axes.plot(x, y, color = "red", linewidth=0.5)

    def genSpeedFigures(self):
        width = self.figsizes['speed_size'][0]
        height = self.figsizes['speed_size'][1]
        plt.rcParams['font.family'] = 'Times New Roman, SimSun'
        plt.rcParams['xtick.direction'] = 'in'
        plt.rcParams['ytick.direction'] = 'in'
        # plt.rcParams['axes.labelsize'] = 14  # 坐标轴标签
        # plt.rcParams['xtick.labelsize'] = 12  # x轴刻度
        # plt.rcParams['ytick.labelsize'] = 12  # y轴刻度
        for load, datalist in self.allData.items():
            for i, csData in enumerate(datalist):
                figure_name = 'ME' + load + ' ' + csData.trialNo + '航速迹程图'
                figure = self.figureManager.get_figure(figure_name)
                figureCanvas = self.figureManager.get_canvas(figure_name)
                if not figure:
                    figure = Figure(figsize=(width, height))
                    figureCanvas = FigureCanvas(figure)
                    self.figureManager.add_figure(csData.seq, figure_name, figure, figureCanvas)
                    csData.speedFigName = figure_name
                ax = figure.add_subplot(111)
                ax_f = ax.twinx()
                line_speed, = ax.plot(csData.df['time'], csData.df['speed'], color='blue', linewidth=2, label='SPEED')
                line_distance, = ax_f.plot(csData.df['time'], csData.df['distance'], color='pink', linewidth=2, label='DISTANCE')
                
                title = 'SPEED & DISTANCE Vs TIME'
                ax.set_xlabel('TIME(s)')
                ax.set_ylabel('SPEED(kn)')
                ax_f.set_ylabel('DISTANCE(m)')
                ax.set_title(title)
                ax.legend(handles=[line_speed, line_distance], loc='lower right')
                # ax_f.legend()

                ax.set_xlim(csData.df['time'].iloc[0], csData.df['time'].iloc[-1])
                ax.set_ylim(int(csData.df['speed'].iloc[0]) - 7, int(csData.df['speed'].iloc[0]) + 3)
                ax_f.set_ylim(csData.df['distance'].iloc[0], int(round(csData.df['distance'].iloc[-1] * 1.25, -2)))
                
                ax.spines['bottom'].set_linewidth(2)
                ax.spines['top'].set_linewidth(2)
                ax.spines['left'].set_linewidth(2)
                ax.spines['right'].set_linewidth(2)
                ax.grid(True)
                figure.tight_layout()

                # 保存图片，调试时打开
                figureCanvas.draw()
                path = f'{figure_name}.png'
                figureCanvas.print_png(path)


    def updateReport(self, index: int, **kwargs):
        self.logger.info(f"update report: {index} {kwargs}")
        for field, value in kwargs.items():
            self.logger.info(f"Field: {field}, Value: {value}")

    def saveReport(self):
        self.reportDoc = Document()
        self.setToA4(self.reportDoc)
        reportFile = f"{self.REPORT_PATH}csReport.doc"
        self.logger.debug("savr cs report to " + reportFile)
        
        # 保存图片
        self.saveFigures()

        # 输出表格
        self.genMutliVoyageTable()

        # 输出轨迹图、航速迹程图到docx中
        for load, datalist in self.allData.items():
            for i, csData in enumerate(datalist):
                self.genSingleVoyageFigure(csData)

        # 保存到docx中
        self.reportDoc.save(reportFile)
        return reportFile

    def saveFigures(self):
        import os, shutil
        tmp_path = 'tmp'
        # 如果tmp文件夹已经存在并且有内容，删除文件夹并重新创建
        if os.path.exists(tmp_path) or os.path.isdir(tmp_path):
            shutil.rmtree(tmp_path)
        os.makedirs(tmp_path)
        for load, datalist in self.allData.items():
            for i, csData in enumerate(datalist):
                trace_path = f'tmp/{csData.traceFigName}.png'
                self.figureManager.save_figure(csData.traceFigName, trace_path)
                speed_path = f'tmp/{csData.speedFigName}.png'
                self.figureManager.save_figure(csData.speedFigName, speed_path)

    def genMutliVoyageTable(self):
        # 读取模板表格
        templatePath = './templates/pow_template2.docx'
        tableTem = Document(templatePath)
        temTable = tableTem.tables[0]

        # 创建目标表格
        cols = 7 + len(self.allData) + self.voyageCount
        powTable = self.reportDoc.add_table(14, cols)
        powTable.style = 'Table Grid'
        powTable.alignment = temTable.alignment
        powTable.autofit = temTable.autofit

        # 根据模板表格修改目标表格格式
        '''参数信息列修改'''
        for i in range(0, 6):
            copy_column_format_and_content(temTable, powTable, i, i)
        '''工况数据列修改'''
        for i in range(6, cols - 1):
            copy_column_format_and_content(temTable, powTable, 6, i)
        '''平均值列修改'''
        meanIndex = []
        start = 6
        for load, datalist in self.allData.items():
            meanIndex.append(start + len(datalist))
            start += len(datalist) + 1
        for i in meanIndex:
            copy_column_format_and_content(temTable, powTable, 8, i)
        '''Note列修改'''
        copy_column_format_and_content(temTable, powTable, 23, cols - 1)

        # 船舶及环境参数
        self.setTableParagraphSpacing(powTable, space_after=0)
        self.replaceCellContent(powTable.columns[0].cells[1], str(self.shipParams.get('airTemperature', '0')))
        self.replaceCellContent(powTable.columns[0].cells[3], self.shipParams.get('weather', 'unknown'))
        self.replaceCellContent(powTable.columns[0].cells[8], str(self.shipParams.get('fwdDraught', '0')))
        self.replaceCellContent(powTable.columns[0].cells[10], self.shipParams.get('vesselName', 'unknown'))
        self.replaceCellContent(powTable.columns[1].cells[1], str(self.shipParams.get('airPressure', '1000')))
        self.replaceCellContent(powTable.columns[1].cells[4], str(self.shipParams.get('waveScale', '0')))
        self.replaceCellContent(powTable.columns[1].cells[8], str(self.shipParams.get('midDraught', '0')))
        self.replaceCellContent(powTable.columns[1].cells[10], self.shipParams.get('hullNo', 'unknown'))
        self.replaceCellContent(powTable.columns[2].cells[1], str(self.shipParams.get('waterTemperature', '0')))
        self.replaceCellContent(powTable.columns[2].cells[3], self.shipParams.get('windDirection', 'NE'))
        self.replaceCellContent(powTable.columns[2].cells[8], str(self.shipParams.get('aftDraught', '0')))
        self.replaceCellContent(powTable.columns[2].cells[10], self.shipParams.get('date', 'unknown'))
        self.replaceCellContent(powTable.columns[3].cells[1], str(self.shipParams.get('waterDensity', '1')))
        self.replaceCellContent(powTable.columns[3].cells[4], str(self.shipParams.get('windScale', '0')))
        self.replaceCellContent(powTable.columns[3].cells[8], str(self.shipParams.get('displacement', 'unknown')))
        self.replaceCellContent(powTable.columns[3].cells[10], self.shipParams.get('seaArea', 'unknown'))

        # 工况数据填表
        start_index = 6
        for load, datalist in self.allData.items():
            for i, data in enumerate(datalist):
                self.fillVoyageData(powTable, start_index + i, data)
            if len(datalist) != 0:
                self.fillMeanValue(powTable, start_index, start_index + len(datalist) - 1)
            start_index += len(datalist) + 1
        
        
        # self.copyTableToDocument(tableTem, self.reportDoc)

        # 合并单元格
        for i in range(0, 4):
            if i%2 == 0:
                mergeCellWithOneValue(powTable.cell(3, i), powTable.cell(4, i))
            mergeCellWithOneValue(powTable.cell(5, i), powTable.cell(6, i))
            mergeCellWithOneValue(powTable.cell(10, i), powTable.cell(12, i))
            # powTable.cell(5, i).merge(powTable.cell(6, i))
            # powTable.cell(10, i).merge(powTable.cell(12, i))
        for i in range(4, len(powTable.rows[0].cells) - 1):
            mergeCellWithOneValue(powTable.cell(4, i), powTable.cell(5, i))
            mergeCellWithOneValue(powTable.cell(6, i), powTable.cell(7, i))
            # powTable.cell(4, i).merge(powTable.cell(5, i))
            # powTable.cell(6, i).merge(powTable.cell(7, i))
            if i == 4:
                mergeCellWithOneValue(powTable.cell(9, i), powTable.cell(10, i))
                # powTable.cell(9, i).merge(powTable.cell(10, i))
        mergeCellWithOneValue(powTable.cell(0, len(powTable.rows[0].cells) - 1), powTable.cell(12, len(powTable.rows[0].cells) - 1))
        # powTable.cell(0, len(powTable.rows[0].cells) - 1).merge(powTable.cell(12, len(powTable.rows[0].cells) - 1))
        for i in range(11, 14):
            mergeCellWithOneValue(powTable.cell(i, 4), powTable.cell(i, 5))
            # powTable.cell(i, 4).merge(powTable.cell(i, 5))
        
        # 插入分页符
        paragraph = self.reportDoc.add_paragraph()
        run = paragraph.add_run()
        run.add_break(WD_BREAK.PAGE)

    def genSingleVoyageFigure(self, data):
        
        table = self.reportDoc.add_table(19, 4)
        table.style = 'Table Grid'

        # 设置表格总宽度为 17 厘米
        set_table_width(table, 17)

        # 设置表格列宽
        cell_width = [5.2, 1.4, 1.1, 9]
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                # print(i, cell_width[i])
                cell.width = Cm(cell_width[i])

        # 合并单元格
        table.cell(0,0).merge(table.cell(0,3))
        table.cell(1,0).merge(table.cell(1,3))
        table.cell(2,0).merge(table.cell(2,2))
        table.cell(3,1).merge(table.cell(3,2))
        table.cell(4,1).merge(table.cell(4,2))
        table.cell(13,0).merge(table.cell(13,2))
        table.cell(17,0).merge(table.cell(17,2))
        table.cell(18,0).merge(table.cell(18,2))
        table.cell(2,3).merge(table.cell(17,3))

        # 设置所有单元格内容格式
        for row in table.rows:
            for cell in row.cells:
                set_cell_content_format(cell, style = WD_PARAGRAPH_ALIGNMENT.CENTER, font_size = Pt(9), font_name = 'Times New Roman')

        # 设置单元格背景色
        set_cell_background(table.cell(2, 0), 'FFFF00') # 单元格黄色背景
        set_cell_background(table.cell(13, 0), 'FFFF00')
        set_cell_background(table.cell(17, 0), 'FFFF00')
        
        # # 填单元格内容
        set_cell_content(table.cell(3, 0), 'START TIME')
        set_cell_content(table.cell(4, 0), 'END TIME')
        set_cell_content(table.cell(5, 0), 'WIND DIRECTION')
        set_cell_content(table.cell(6, 0), 'WIND SPEED')
        set_cell_content(table.cell(7, 0), 'INITIAL HEADING')
        set_cell_content(table.cell(8, 0), 'INITIAL SPEED')
        set_cell_content(table.cell(9, 0), 'FINAL HEADING')
        set_cell_content(table.cell(10, 0), 'FINAL SPEED')
        set_cell_content(table.cell(11, 0), 'TOTAL TIME TAKEN')
        set_cell_content(table.cell(12, 0), 'DISTANCE')
        set_cell_content(table.cell(14, 0), 'MEASUREMENT SPEED')
        set_cell_content(table.cell(15, 0), 'MEASUREMENT ROTATE SPEED')
        set_cell_content(table.cell(16, 0), 'MEASUREMENT SHAFT POWER')

        set_cell_content(table.cell(5, 2), 'deg')
        set_cell_content(table.cell(6, 2), 'm/s')
        set_cell_content(table.cell(7, 2), 'deg')
        set_cell_content(table.cell(8, 2), 'kn')
        set_cell_content(table.cell(9, 2), 'deg')
        set_cell_content(table.cell(10, 2), 'kn')
        set_cell_content(table.cell(11, 2), 'sec')
        set_cell_content(table.cell(12, 2), 'm')
        set_cell_content(table.cell(14, 2), 'kn')
        set_cell_content(table.cell(15, 2), 'r/min')
        set_cell_content(table.cell(16, 2), 'kW')

        set_cell_content(table.cell(3, 1), data.startTime)
        set_cell_content(table.cell(4, 1), data.endTime)
        set_cell_content(table.cell(5, 1), data.windDirection)
        set_cell_content(table.cell(6, 1), data.windSpeed)
        set_cell_content(table.cell(7, 1), data.initialHeading)
        set_cell_content(table.cell(8, 1), data.initialSpeed)
        set_cell_content(table.cell(9, 1), data.finalHeading)
        set_cell_content(table.cell(10, 1), data.finalSpeed)
        set_cell_content(table.cell(11, 1), data.period)
        set_cell_content(table.cell(12, 1), data.distance)
        set_cell_content(table.cell(14, 1), data.speed)
        set_cell_content(table.cell(15, 1), data.rpm)
        set_cell_content(table.cell(16, 1), data.shaftPower)

        table_title = f'     Celsius Essen                         SPEED TRIAL AT M/E {data.rpm} r/min'
        set_cell_content_format(table.cell(0, 0), table_title, font_bold=True, font_size=Pt(14), font_color=RGBColor(231,27,100))

        dates = data.startTime.split(' ')
        date = 'Date: ' + dates[0]
        set_cell_content_format(table.cell(18, 0), date, style=WD_PARAGRAPH_ALIGNMENT.LEFT)

        trial = 'TRIAL: ' + data.trialNo
        set_cell_content_format(table.cell(18, 3), trial, style=WD_PARAGRAPH_ALIGNMENT.RIGHT)

        # 向表格中插入图片
        path = f'tmp/{data.traceFigName}.png'
        add_image_to_cell(table, 1, 0, path, Cm(15), Cm(15))
        path1 = f'tmp/{data.speedFigName}.png'
        add_image_to_cell(table, 2, 3, path1, Cm(8.7), Cm(6.5))

        # 插入分页符
        paragraph = self.reportDoc.add_paragraph()
        run = paragraph.add_run()
        run.add_break(WD_BREAK.PAGE)

    def fillVoyageData(self, powTable, index, data):
        self.replaceCellContent(powTable.columns[index].cells[0], str(data.speed))
        self.replaceCellContent(powTable.columns[index].cells[1], str(data.distance))
        self.replaceCellContent(powTable.columns[index].cells[2], str(data.period))
        self.replaceCellContent(powTable.columns[index].cells[3], str(data.heading))
        self.replaceCellContent(powTable.columns[index].cells[4], str(data.shaftPower))
        self.replaceCellContent(powTable.columns[index].cells[6], str(data.rpm))
        self.replaceCellContent(powTable.columns[index].cells[8], str(data.depth))
        self.replaceCellContent(powTable.columns[index].cells[9], str(data.windSpeed))
        self.replaceCellContent(powTable.columns[index].cells[10], str(data.windDirection))
        self.replaceCellContent(powTable.columns[index].cells[11], str(data.MELoad))
        self.replaceCellContent(powTable.columns[index].cells[12], str(data.startTime))
        self.replaceCellContent(powTable.columns[index].cells[13], str(data.trialNo))
    
    def fillMeanValue(self,powTable, start, end):
        num = end - start + 1
        sumRpm = 0
        sumPower = 0
        sumSpeed = 0
        for i in range(start, end + 1):
            sumRpm += int(powTable.columns[i].cells[6].text)
            sumPower += float(powTable.columns[i].cells[4].text)
            sumSpeed += float(powTable.columns[i].cells[0].text)
        meanRpm = round(sumRpm * 1.0 / num)
        meanPower = round(sumPower / num, 1)
        meanSpeed = round(sumSpeed / num, 2)
        self.replaceCellContent(powTable.columns[end + 1].cells[0], str(meanRpm))
        self.replaceCellContent(powTable.columns[end + 1].cells[4], str(meanPower))
        self.replaceCellContent(powTable.columns[end + 1].cells[6], str(meanSpeed))

        self.replaceCellContent(powTable.columns[end + 1].cells[1], '---')
        self.replaceCellContent(powTable.columns[end + 1].cells[2], '---')
        self.replaceCellContent(powTable.columns[end + 1].cells[3], '---')
        self.replaceCellContent(powTable.columns[end + 1].cells[8], '---')
        self.replaceCellContent(powTable.columns[end + 1].cells[9], '---')
        self.replaceCellContent(powTable.columns[end + 1].cells[10], '---')
        self.replaceCellContent(powTable.columns[end + 1].cells[11], '---')
        self.replaceCellContent(powTable.columns[end + 1].cells[12], '---')
        self.replaceCellContent(powTable.columns[end + 1].cells[13], 'Mean')

    def setToA4(self, doc):
        """
        设置文档页面大小为 A4 (210mm × 297mm)
        :param doc: Document 对象
        """
        section = doc.sections[0]
        section.page_width = Mm(210)  # A4 宽度
        section.page_height = Mm(297)  # A4 高度

    def copyTableToDocument(self, source_doc, target_doc, table_index=0):

        # 提取指定表格
        if table_index >= len(source_doc.tables):
            raise IndexError("表格索引超出范围")
        
        table_to_copy = source_doc.tables[table_index]

        # 复制表格的 XML 元素
        table_xml = deepcopy(table_to_copy._element)

        # 将表格插入到目标文档末尾
        target_doc._body._element.append(table_xml)

        # 保存目标文档
        #target_doc.save(save_path)
        #print(f"表格已成功复制到文档并保存到：{save_path}")

    def replaceCellContent(self, cell, new_text):
        """
        替换表格单元格内容，同时保持原有格式。
        :param cell: 要修改的单元格对象
        :param new_text: 替换后的文本内容
        """
        # 清空单元格的所有段落内容
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.text = ""

        # 将新文本写入第一个段落，保留格式
        if cell.paragraphs:
            paragraph = cell.paragraphs[0]
            if not paragraph.runs:  # 如果 runs 为空
                new_run = paragraph.add_run(new_text)  # 添加一个新的 Run
                # new_run.bold = True  # 设置字体加粗
                # new_run.italic = True  # 设置字体斜体
                # print("成功添加新文本到空段落中")
            else:
                paragraph.runs[0].text = new_text  # 写入第一个运行
        else:
            # 如果单元格为空，则创建一个段落和运行
            paragraph = cell.add_paragraph()
            paragraph.add_run(new_text)

    def setTableParagraphSpacing(self, table, space_before=None, space_after=None, line_spacing=None):
        """
        设置表格中所有单元格的段落间距
        :param table: 表格对象
        :param space_before: 段前间距（单位：磅）
        :param space_after: 段后间距（单位：磅）
        :param line_spacing: 行间距（支持倍数或 None）
        """
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph_format = paragraph.paragraph_format
                    if space_before is not None:
                        paragraph_format.space_before = Pt(space_before)  # 设置段前间距
                    if space_after is not None:
                        paragraph_format.space_after = Pt(space_after)  # 设置段后间距
                    if line_spacing is not None:
                        paragraph_format.line_spacing = line_spacing  # 设置行间距

if __name__ == '__main__':
    csr = CSReportGenImpl()
    reportFile = f"{csr.REPORT_PATH}csReport.doc"

    csr.reportDoc = Document()
    csr.setToA4(csr.reportDoc)
    tableTem = Document('./templates/pow_template.docx')
    powTable = tableTem.tables[0]
    print(powTable.columns[0].cells[1].text)
    csr.setTableParagraphSpacing(powTable, space_after=0)
    csr.replaceCellContent(powTable.columns[0].cells[1], str(3))

    i = 0
    for c in powTable.columns[0].cells:
        print(i, c.text)
        i+=1
    #csr.copyTableToDocument(tableTem, csr.reportDoc)
    target_table = csr.reportDoc.add_table(14, 24)
    target_table.style = 'Table Grid'
    target_table.alignment = powTable.alignment
    target_table.autofit = powTable.autofit
    # copy_column_format_and_content(powTable, target_table, 0, 0)
    # copy_column_format_and_content(powTable, target_table, 5, 5)
    for i in range(0, 6):
        copy_column_format_and_content(powTable, target_table, i, i)
    for i in range(6, 8):
        copy_column_format_and_content(powTable, target_table, 6, i)
    copy_column_format_and_content(powTable, target_table, 8, 8)
    copy_column_format_and_content(powTable, target_table, 23, 9)

    for i in range(0, 4):
        if i%2 == 0:
            mergeCellWithOneValue(target_table.cell(3, i), target_table.cell(4, i))
        mergeCellWithOneValue(target_table.cell(5, i), target_table.cell(6, i))
        mergeCellWithOneValue(target_table.cell(10, i), target_table.cell(12, i))
    for i in range(4, len(target_table.rows[0].cells) - 1):
        mergeCellWithOneValue(target_table.cell(4, i), target_table.cell(5, i))
        mergeCellWithOneValue(target_table.cell(6, i), target_table.cell(7, i))
        if i == 4:
            mergeCellWithOneValue(target_table.cell(9, i), target_table.cell(10, i))
    mergeCellWithOneValue(target_table.cell(0, 9), target_table.cell(12, 9))
    csr.reportDoc.save(reportFile)
